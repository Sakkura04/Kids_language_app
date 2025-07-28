import pdfplumber
from docx import Document
import os
import re

def clean_line(line):
    return ''.join(c for c in line if c >= ' ' or c in '\t\n\r')

def is_unwanted_line(line):
    # Skip lines with .indd (file names)
    if '.indd' in line:
        return True
    # Skip lines that look like dates/times (very basic pattern)
    if re.search(r'\d{1,2}/\d{1,2}/\d{2,4}', line):
        return True
    if re.search(r'\d{1,2}:\d{2}:\d{2}', line):
        return True
    # Skip lines that are just a page number (digits only, possibly with whitespace)
    if re.fullmatch(r'\s*\d+\s*', line):
        return True
    return False

def merge_lines(lines):
    # Merge lines unless the sentence ends with a dot and is followed by two line breaks
    merged = []
    buffer = ''
    i = 0
    while i < len(lines):
        line = lines[i]
        if is_unwanted_line(line):
            i += 1
            continue
        line = clean_line(line)
        if not line:
            # Check for double line break after a dot
            if merged and merged[-1].endswith('.') and i+1 < len(lines) and not lines[i+1].strip():
                merged.append('')  # keep the blank line
                i += 1
                continue
            # Otherwise, skip single blank lines
            i += 1
            continue
        if buffer:
            buffer += ' ' + line
        else:
            buffer = line
        # If this is the last line or next line is blank, flush buffer
        if i+1 == len(lines) or not lines[i+1].strip():
            merged.append(buffer)
            buffer = ''
        i += 1
    if buffer:
        merged.append(buffer)
    return merged

def pdf_to_word(pdf_path, docx_path, num_columns=2, start_page=0, end_page=None):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        if end_page is None or end_page >= total_pages:
            end_page = total_pages - 1
        for i, page in enumerate(pdf.pages):
            if i < start_page:
                continue
            if i > end_page:
                break
            width = page.width
            height = page.height
            col_width = width / num_columns
            for col in range(num_columns):
                left = col * col_width
                right = (col + 1) * col_width
                bbox = (left, 0, right, height)
                col_page = page.within_bbox(bbox)
                col_text = col_page.extract_text()
                if col_text:
                    lines = col_text.split('\n')
                    merged_lines = merge_lines(lines)
                    for line in merged_lines:
                        if line != '':
                            doc.add_paragraph(line)
                        else:
                            doc.add_paragraph('')
    try:
        doc.save(docx_path)
        print(f"Saved Word document to: {docx_path}")
    except PermissionError:
        print(f"ALARM: Could not save the Word document. Please CLOSE '{os.path.basename(docx_path)}' if it is open, then run the program again.")

if __name__ == "__main__":
    # List all PDF files in the current directory
    current_dir = os.path.dirname(__file__)
    pdf_files = [f for f in os.listdir(current_dir) if f.lower().endswith('.pdf')]
    if not pdf_files:
        print("No PDF files found in the current directory.")
        exit(1)
    print("Available PDF files:")
    for f in pdf_files:
        print(f"- {f}")
    pdf_name = input("Enter the exact PDF file name to parse (with .pdf extension): ").strip()
    while pdf_name not in pdf_files:
        print("File not found. Please enter a valid file name from the list above.")
        pdf_name = input("Enter the exact PDF file name to parse (with .pdf extension): ").strip()
    pdf_path = os.path.join(current_dir, pdf_name)
    docx_name = os.path.splitext(pdf_name)[0] + '.docx'
    docx_path = os.path.join(current_dir, docx_name)
    try:
        num_columns = int(input("How many columns does the text have per page? (default 2): ") or 2)
        start_page = int(input("From which page number should parsing begin? (0-based, default 0): ") or 0)
        end_page = input("At which page number should parsing stop? (inclusive, 0-based, leave blank for last page): ")
        end_page = int(end_page) if end_page.strip() != '' else None
    except ValueError:
        print("Invalid input. Using defaults: 2 columns, start from page 0, parse to last page.")
        num_columns = 2
        start_page = 0
        end_page = None
    pdf_to_word(pdf_path, docx_path, num_columns=num_columns, start_page=start_page, end_page=end_page) 